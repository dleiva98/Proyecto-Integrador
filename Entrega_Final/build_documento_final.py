# -*- coding: utf-8 -*-
"""
Genera el Documento Final de Entrega del Proyecto Integrador (Equipo 66).
Consolida las Entregas 1-2 y los Avances 3-6 en un informe academico unico,
con portada institucional, resumen ejecutivo/abstract, indice automatico,
tablas, figuras y referencias en formato APA.

Fuentes de contenido: README.md del repositorio y Avance 6 (material propio
del equipo). Salida: Documento_Final_Entrega_Equipo66.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/home/user/Proyecto-Integrador"
ASSETS = os.path.join(ROOT, "Entrega_Final", "assets")
OUT = os.path.join(ROOT, "Entrega_Final", "Documento_Final_Entrega_Equipo66.docx")

AZUL = RGBColor(0x00, 0x39, 0xA6)       # Azul Tec
AZUL_OSC = RGBColor(0x0C, 0x23, 0x40)   # Azul marino institucional
GRIS = RGBColor(0x5A, 0x60, 0x68)
GRIS_TAB = "DCE3F0"                       # encabezado de tabla (azul muy claro)
NEGRO = RGBColor(0x20, 0x20, 0x20)
BODY_FONT = "Calibri"
SERIF = "Cambria"

doc = Document()

# ---------------------------------------------------------------- estilos base
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = NEGRO
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(8)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, sz, col in [("Heading 1", 17, AZUL), ("Heading 2", 13.5, AZUL_OSC),
                       ("Heading 3", 12, AZUL_OSC)]:
    st = doc.styles[name]
    st.font.name = SERIF
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = col
    st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

# márgenes
for s in doc.sections:
    s.top_margin = Cm(2.4); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.6); s.right_margin = Cm(2.6)


# ------------------------------------------------------------------- utilidades
def set_run(r, size=None, bold=None, italic=None, color=None, font=None):
    if size is not None: r.font.size = Pt(size)
    if bold is not None: r.font.bold = bold
    if italic is not None: r.font.italic = italic
    if color is not None: r.font.color.rgb = color
    if font is not None: r.font.name = font
    return r


def para(text="", style=None, align=None, size=None, bold=None, italic=None,
         color=None, font=None, space_after=None, space_before=None):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        set_run(r, size, bold, italic, color, font)
    return p


def rich(parts, style=None, align=None, space_after=None):
    """parts: lista de (texto, dict-de-formato)."""
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    for txt, fmt in parts:
        set_run(p.add_run(txt), **fmt)
    return p


def shade(cell, hexc):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(sh)


def set_cell(cell, text, bold=False, color=None, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT,
             white=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = BODY_FONT
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        r.font.color.rgb = color


def table(headers, rows, widths=None, header_fill=GRIS_TAB, aligns=None,
          highlight_row=None, font_size=9.5, caption=None):
    if caption:
        cp = para(caption, size=9.5, italic=True, color=GRIS, space_after=4,
                  space_before=2)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=AZUL_OSC, size=font_size,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            a = WD_ALIGN_PARAGRAPH.LEFT if aligns is None else aligns[ci]
            set_cell(cells[ci], str(val), size=font_size, align=a)
            if highlight_row is not None and ri == highlight_row:
                shade(cells[ci], "EAF0FB")
                cells[ci].paragraphs[0].runs[0].font.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(path, caption, width_in=6.3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(path, width=Inches(width_in))
    cp = para(caption, size=9, italic=True, color=GRIS,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    return cp


def bullets(items, style="List Bullet"):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style=style)
            set_run(p.add_run(it[0]), bold=True)
            set_run(p.add_run(it[1]))
        else:
            doc.add_paragraph(it, style=style)


def hrule(color="0039A6", size=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)
    return p


def page_break():
    doc.add_page_break()


# ============================================================== PORTADA
sec = doc.sections[0]
para(space_after=2)
# logo
lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.add_run().add_picture(os.path.join(ASSETS, "logo_tec.png"), width=Inches(3.5))
lp.paragraph_format.space_after = Pt(2)

para("Escuela de Ingeniería y Ciencias", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=12, color=GRIS, space_after=0)
para("Maestría en Inteligencia Artificial Aplicada",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12.5, bold=True, color=AZUL_OSC,
     space_after=2)
hrule()
para(space_after=14)

para("PROYECTO INTEGRADOR · DOCUMENTO FINAL DE ENTREGA",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=AZUL,
     space_after=16)

para("Traducción Automática Neuronal para la Preservación de la Lengua "
     "Bribri en Riesgo de Extinción",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=21, bold=True, color=AZUL_OSC,
     font=SERIF, space_after=6)
para("Desarrollo de un Sistema NMT Bribri–Español mediante Arquitecturas "
     "Transformer para Lenguas de Bajos Recursos",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=13.5, italic=True, color=GRIS,
     font=SERIF, space_after=18)

para(space_after=10)
para("Programa Voces — Componente B", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=12.5, bold=True, color=AZUL, space_after=2)
para("Centro Nacional de Inteligencia Artificial (CENIA, Chile) · "
     "Pontificia Universidad Católica de Chile · Tecnológico de Monterrey",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=GRIS, space_after=22)

# bloque de autoría centrado
auth = [
    ("Equipo 66 — Costa Rica", {"size": 12, "bold": True, "color": AZUL_OSC}),
]
rich(auth, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Daniel Leiva  (A01795876)   ·   Israel Agustín Vargas Monroy  (A01796556)",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=10)
rich([("Asesor:  ", {"bold": True, "color": AZUL_OSC}),
      ("Carlos Villaseñor", {})], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
rich([("Patrocinador:  ", {"bold": True, "color": AZUL_OSC}),
      ("Carlos Aspillaga — Director del Programa Voces (CENIA)", {})],
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
rich([("Metodología:  ", {"bold": True, "color": AZUL_OSC}),
      ("CRISP-ML(Q) — Fase de Evaluación y Despliegue", {})],
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

hrule()
para("Suretka, Talamanca, Costa Rica  ·  Santiago, Chile  ·  Junio de 2026",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=GRIS, space_after=0)

page_break()

# ============================================================== RESUMEN
para("Resumen ejecutivo", style="Heading 1")
para(
    "Este documento constituye la entrega final del Proyecto Integrador de la "
    "Maestría en Inteligencia Artificial Aplicada del Tecnológico de Monterrey, "
    "desarrollado por el Equipo 66 (Costa Rica) en el marco del Programa Voces, "
    "iniciativa de preservación de lenguas indígenas coordinada por el Centro "
    "Nacional de Inteligencia Artificial (CENIA, Chile) en colaboración con la "
    "Pontificia Universidad Católica de Chile y el Tecnológico de Monterrey. El "
    "trabajo aborda el Componente B del proyecto: la construcción y el ajuste "
    "fino (fine-tuning) de modelos de Traducción Automática Neuronal (NMT) para "
    "el par bribri–español, lengua de la familia chibcha hablada en el "
    "Territorio Indígena Bribri de Talamanca y clasificada en riesgo.")
para(
    "A partir de un corpus paralelo de 1 505 pares construido desde fuentes "
    "documentales heterogéneas, se realizó fine-tuning de NLLB-200-distilled-600M "
    "y, como baseline comparativo, de M2M-100. El mejor modelo —NLLB-200 "
    "reentrenado en servidor H100 (8 épocas, lr 2e-4)— alcanza chrF 31.43 / "
    "chrF++ 30.47 / spBLEU 21.16 sobre el conjunto de prueba, superando al "
    "baseline M2M-100 por un factor de 3.3× en chrF. Frente a los criterios de "
    "éxito formalizados en la fase de evaluación, el sistema no es apto para "
    "traducción autónoma de propósito general, pero sí cumple como prueba de "
    "concepto funcional para asistencia documental, lexicográfica y pedagógica "
    "con validación humana en el bucle (human-in-the-loop).")
para(
    "La decisión fundamentada es implementar en producción restringida e iterar "
    "en paralelo, sin retroceder a fases anteriores: el modelado cumplió su "
    "objetivo y la limitación del corpus es una mejora planificada, no un defecto "
    "a corregir. Se entrega a CENIA una base reproducible, trazable y medible, "
    "con palancas de mejora priorizadas, un conjunto de accionables asignados por "
    "stakeholder y una arquitectura de despliegue recomendada sobre Microsoft "
    "Azure, por continuidad con la infraestructura de cómputo del proyecto.")
rich([("Palabras clave:  ", {"bold": True, "color": AZUL_OSC}),
      ("traducción automática neuronal; lenguas de bajos recursos; bribri; "
       "NLLB-200; fine-tuning; preservación lingüística; CRISP-ML(Q); chrF.",
       {"italic": True})], space_after=10)

para("Abstract", style="Heading 2")
para(
    "This document is the final deliverable of the Capstone Project for the "
    "Master's in Applied Artificial Intelligence at Tecnológico de Monterrey, "
    "carried out by Team 66 (Costa Rica) within the Voces Program, a "
    "low-resource language preservation initiative led by Chile's National "
    "Center for Artificial Intelligence (CENIA) together with PUC Chile and "
    "Tecnológico de Monterrey. The work addresses Component B: building a "
    "Bribri–Spanish parallel corpus and fine-tuning Neural Machine Translation "
    "models for this endangered Chibchan language. From 1,505 parallel pairs, "
    "NLLB-200-distilled-600M was fine-tuned (with M2M-100 as a comparative "
    "baseline); the best model reaches chrF 31.43 / chrF++ 30.47 / spBLEU 21.16 "
    "on the test set, outperforming the M2M-100 baseline 3.3× in chrF. Measured "
    "against success criteria appropriate to a low-resource preservation goal, "
    "the system qualifies as a functional, human-in-the-loop proof of concept "
    "for documentary and pedagogical assistance, and is recommended for a "
    "restricted production deployment with a parallel improvement roadmap.",
    italic=True)

page_break()

# ============================================================== ÍNDICE
para("Contenido", style="Heading 1")
# campo TOC (Word lo actualiza con F9)
tp = doc.add_paragraph()
run = tp.add_run()
fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
ftxt = OxmlElement("w:r"); ftt = OxmlElement("w:t")
ftt.text = "Actualice este campo en Word (clic derecho ▸ Actualizar campos / F9)."
ftxt.append(ftt)
fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
run._r.append(fb); run._r.append(instr); run._r.append(fs)
run._r.append(ftxt); run._r.append(fe)

page_break()

# ============================================================== 1. INTRODUCCIÓN
para("1. Introducción", style="Heading 1")

para("1.1 Contexto y motivación", style="Heading 2")
para(
    "El bribri es una lengua de la familia chibcha hablada en el Territorio "
    "Indígena Bribri de Talamanca, en la región fronteriza entre Costa Rica y "
    "Panamá, territorio declarado Reserva de la Biosfera La Amistad por la UNESCO "
    "en 1982. Según el Censo Nacional 2011 (INEC, 2011) la lengua es hablada por "
    "aproximadamente 12 000 personas, aunque estudios recientes documentan una "
    "reducción de hablantes activos, particularmente entre las generaciones "
    "jóvenes (Jara Murillo, 2018). Esta erosión intergeneracional sitúa al bribri "
    "en una categoría de riesgo y motiva intervenciones de revitalización "
    "asistida por tecnología.")
para(
    "La ausencia de herramientas digitales de traducción automática para el par "
    "bribri–español limita la transmisión intergeneracional de la lengua, "
    "restringe el acceso de la comunidad a contenidos en su idioma y reduce las "
    "posibilidades de documentación y revitalización lingüística (Gómez-Rendón, "
    "2017; Mager et al., 2018). El presente proyecto responde a esa carencia "
    "construyendo, por primera vez para este par, un sistema NMT reproducible y "
    "medible.")

para("1.2 Problema y justificación", style="Heading 2")
para(
    "El bribri es una lengua de bajos recursos (low-resource language): a "
    "diferencia de lenguas con amplia presencia digital, cuenta con un corpus "
    "paralelo extremadamente limitado, disperso en diccionarios bilingües, "
    "materiales del Ministerio de Educación Pública de Costa Rica (MEP) y "
    "documentación lingüística académica. Esta escasez de datos es la principal "
    "dificultad técnica del proyecto y exige estrategias específicas —"
    "transferencia de aprendizaje, ajuste fino de modelos multilingües "
    "preentrenados y aumento de datos— que sitúan el trabajo en la frontera de "
    "la investigación aplicada en PLN para lenguas indígenas latinoamericanas "
    "(Mager et al., 2018; Ortega et al., 2020).")

para("1.3 Objetivos", style="Heading 2")
rich([("Objetivo general.  ", {"bold": True, "color": AZUL_OSC}),
      ("Desarrollar y evaluar un sistema de Traducción Automática Neuronal "
       "bribri–español, basado en arquitecturas Transformer preentrenadas y "
       "ajustadas finamente sobre un corpus paralelo propio, que sirva como "
       "prueba de concepto reproducible para la preservación y revitalización "
       "lingüística del bribri en el marco del Programa Voces.", {})])
para("Objetivos específicos:", bold=True, color=AZUL_OSC, space_after=2)
bullets([
    "Construir un corpus paralelo bribri–español trazable a partir de fuentes "
    "documentales, con normalización lingüísticamente informada, filtrado de "
    "calidad y partición estratificada reproducible.",
    "Realizar fine-tuning de modelos NMT multilingües (NLLB-200, M2M-100) y "
    "comparar arquitecturas, configuraciones e hiperparámetros bajo un protocolo "
    "de evaluación común.",
    "Evaluar el desempeño con métricas robustas a morfología rica (chrF, "
    "chrF++, spBLEU) y con análisis cualitativo de patrones de error.",
    "Explorar estrategias de ensamble apropiadas para generación de secuencias.",
    "Determinar la aptitud del modelo para despliegue, definir accionables por "
    "stakeholder y proponer una arquitectura de producción.",
])

para("1.4 Alcance — Componente B", style="Heading 2")
para(
    "El proyecto cubre el Componente B del Programa Voces: la contribución "
    "académica de fine-tuning de modelos NMT preentrenados sobre un corpus "
    "paralelo bribri–español. Quedan fuera del alcance la traducción autónoma de "
    "propósito general, el despliegue a escala comercial y el reconocimiento "
    "óptico de caracteres (OCR) de fuentes escaneadas. La tabla siguiente resume "
    "los datos generales del proyecto.")
table(
    ["Dato general", "Detalle"],
    [["Institución / Patrocinador",
      "CENIA (Chile) — Carlos Aspillaga, Director del Programa Voces"],
     ["Par lingüístico", "bribri ↔ español  (bri / es)"],
     ["Dominio", "Procesamiento de Lenguaje Natural — NMT supervisada (Transformer seq2seq)"],
     ["Lugar de aplicación", "Suretka, Talamanca, Limón, Costa Rica"],
     ["Equipo CR", "Daniel Leiva (A01795876); Israel Agustín Vargas Monroy (A01796556)"],
     ["Clasificación SCIAN", "541720 — Investigación y desarrollo en ciencias naturales y exactas"],
     ["Mejor modelo", "NLLB-200-distilled-600M · 8 épocas · lr 2e-4  (chrF 31.43 / spBLEU 21.16)"]],
    widths=[5.0, 11.5])

page_break()

# ============================================================== 2. CONTEXTO
para("2. Contexto institucional: el Programa Voces", style="Heading 1")
para(
    "El Centro Nacional de Inteligencia Artificial (CENIA) es el principal centro "
    "de investigación y desarrollo en IA de Chile, fundado en 2021 en el marco "
    "del programa de centros de excelencia del gobierno chileno. CENIA lidera el "
    "Programa Voces en colaboración con la Pontificia Universidad Católica de "
    "Chile y el Tecnológico de Monterrey, actuando como entidad coordinadora de "
    "los recursos computacionales, la infraestructura de entrenamiento y la "
    "dirección científica.")
para(
    "El trabajo de campo se concentra en el Territorio Indígena Bribri de "
    "Talamanca. Suretka es una de las comunidades de mayor densidad poblacional "
    "bribri y el principal nodo de coordinación con la Asociación de Desarrollo "
    "Integral del Territorio Indígena Bribri de Talamanca (ADITIBRI), entidad con "
    "la que el proyecto mantiene vínculos de colaboración para la validación "
    "cultural y lingüística. Desde Costa Rica, Daniel Leiva actúa como encargado "
    "regional, coordinando la recopilación de datos, las alianzas con entidades "
    "académicas locales —entre ellas la Universidad de Costa Rica— y la "
    "articulación con las comunidades bajo principios de soberanía lingüística y "
    "beneficio comunitario.")

para("3. Marco metodológico: CRISP-ML(Q)", style="Heading 1")
para(
    "El proyecto se estructura según CRISP-ML(Q) (Studer et al., 2021), una "
    "extensión de CRISP-DM con aseguramiento de calidad para sistemas de "
    "aprendizaje automático. El presente documento integra todas sus fases y "
    "culmina en la de Evaluación y Despliegue. La correspondencia con las "
    "entregas del proyecto es la siguiente:")
table(
    ["Fase CRISP-ML(Q)", "Entrega / Avance", "Contenido principal"],
    [["Comprensión del negocio y de los datos", "Entrega 1", "Objetivo de negocio, sponsor, lugar y dominio de aplicación"],
     ["Preparación de datos", "Entrega 2", "Construcción del corpus, normalización, filtros y partición"],
     ["Modelado", "Avances 3–4", "Fine-tuning NLLB-200 (Colab T4 y H100) y baseline M2M-100"],
     ["Modelado — ensambles", "Avance 5", "Checkpoint averaging y system combination (MBR)"],
     ["Evaluación y despliegue", "Avance 6", "Criterios de éxito, decisión, accionables y arquitectura cloud"]],
    widths=[5.2, 3.0, 8.3])

page_break()

# ============================================================== 4. ENTREGA 1
para("4. Comprensión del negocio y de los datos", style="Heading 1")

para("4.1 Título y ejes del proyecto", style="Heading 2")
para(
    "El título —«Traducción Automática Neuronal para la Preservación de la "
    "Lengua Bribri en Riesgo de Extinción»— refleja los tres ejes del proyecto: "
    "la metodología técnica (NMT basada en Transformers), el par lingüístico "
    "(bribri–español) y el propósito institucional y social (preservación de una "
    "lengua indígena costarricense en riesgo). La inclusión del término "
    "low-resource languages responde a la categorización internacional del "
    "bribri en la literatura de PLN.")

para("4.2 Lugar de aplicación", style="Heading 2")
table(
    ["Nivel", "Detalle"],
    [["País", "Costa Rica"],
     ["Provincia / Cantón", "Limón / Talamanca"],
     ["Comunidad", "Suretka, Territorio Indígena Bribri de Talamanca"],
     ["Sede institucional", "CENIA, Santiago, Chile (coordinación central)"]],
    widths=[4.5, 12.0])

para("4.3 Dominio de aplicación", style="Heading 2")
para(
    "El dominio es el Procesamiento de Lenguaje Natural; la técnica principal es "
    "el aprendizaje supervisado mediante Traducción Automática Neuronal sobre una "
    "arquitectura Transformer (seq2seq con mecanismo de atención). La complejidad "
    "reside en la naturaleza de bajos recursos del bribri: para mitigar la "
    "escasez de datos, el proyecto explora transferencia de aprendizaje, ajuste "
    "fino de modelos preentrenados en lenguas tipológicamente relacionadas y "
    "aumento de datos.")

# ============================================================== 5. ENTREGA 2
para("5. Preparación de datos y construcción del corpus", style="Heading 1")
para(
    "Esta fase integra la construcción del corpus, la ingeniería de "
    "características y la partición de datos. La rúbrica de ingeniería de "
    "características (FE) está formulada para datos tabulares (binning, "
    "codificación, escalamiento, PCA); en NMT sobre texto paralelo esas "
    "operaciones no tienen análogo directo y se materializan en operaciones "
    "equivalentes sobre el texto, que cumplen la misma función: convertir datos "
    "crudos en entradas útiles y de baja varianza para el modelo.")

table(
    ["Operación FE tabular", "Equivalente en el pipeline NMT", "Justificación"],
    [["Limpieza / normalización de escala", "Normalización Unicode NFC + reglas por idioma",
      "Tokens idénticos no se representan con bytes distintos; preserva diacríticos del bribri (tonos, nasalización, corte glotal), fonémicamente distintivos."],
     ["Filtrado de outliers / umbral de varianza", "Filtros de longitud, ratio y deduplicación",
      "Descarta pares degenerados (vacíos, longitud extrema, ratio anómalo, duplicados) que introducirían ruido o memorización."],
     ["Codificación de categorías", "Tokenización SentencePiece (subword)",
      "Convierte texto en IDs de subpalabra; el token de idioma condiciona la dirección de traducción."],
     ["Selección de características", "Etiquetado por dominio y confianza",
      "Permite estratificar y filtrar por calidad de fuente sin perder trazabilidad."],
     ["Extracción de características", "Embeddings del Transformer preentrenado",
      "El modelo proyecta cada token a un espacio latente multilingüe; el fine-tuning reajusta esa representación al par bribri–español."]],
    widths=[3.8, 4.5, 8.2], font_size=9)

para("5.1 Esquema de un par paralelo y reglas de normalización", style="Heading 2")
para(
    "Cada par se almacena con metadatos de trazabilidad (documento y página de "
    "origen, dominio, método de extracción y confianza). Los dominios son "
    "didactico, narrativo, religioso, etnografico y web; los niveles de confianza "
    "son high, medium y low. Las reglas de normalización son deliberadamente "
    "asimétricas entre idiomas:")
bullets([
    ("Todo texto:  ", "normalización Unicode NFC."),
    ("Bribri:  ", "preserva todos los diacríticos (tonos, virgulilla nasal, "
     "corte glotal); nunca se convierte a minúsculas ni se eliminan combining marks."),
    ("Español:  ", "NFC + comillas tipográficas a rectas, guiones largos a «-» y "
     "espacios en blanco colapsados."),
])
para("Los filtros se aplican en orden y quedan registrados en el reporte del "
     "pipeline:", space_after=2)
bullets([
    "Descarta si el lado bribri o español está vacío.",
    "Descarta si alguno tiene menos de 2 o más de 80 palabras.",
    "Descarta si el ratio de longitudes (caracteres) está fuera de [0.3, 3.0].",
    "Descarta si ambos lados son idénticos.",
    "Deduplica por SHA-256 del par (bri, es).",
], style="List Number")

para("5.2 Fuentes y resultado del corpus (v0)", style="Heading 2")
para(
    "Tras los filtros, el corpus efectivo es de 1 505 pares (de 1 707 brutos), en "
    "el extremo bajo del rango estimado en el Avance 1, confirmando que el bribri "
    "es un escenario de muy bajo recurso incluso tras agotar las fuentes "
    "documentales disponibles.")
table(
    ["Fuente", "Brutos", "Conservados", "Dominio", "Confianza"],
    [["ESSJ_Sanchez_Avendano_Vol1.pdf", "814", "675", "religioso", "medium"],
     ["I_tte_Historias_bribris.pdf", "774", "728", "narrativo", "low ¹"],
     ["Gramatica_lengua_bribri_Jara2018.pdf", "76", "65", "didactico", "high / medium"],
     ["Sevtto_bribri_ie_Hablemos.pdf", "21", "18", "didactico", "high"],
     ["Palabras_Francisco_Garcia.pdf", "18", "17", "etnografico", "low ¹"],
     ["Ditso_rukuo_Identidad_semillas.pdf", "4", "2", "narrativo", "low"],
     ["https://www.lenguabribri.com/", "0", "0", "web", "—"]],
    widths=[6.5, 1.8, 2.4, 3.0, 2.8], font_size=9,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT])
para(
    "¹ El «español» en I ttè (interlineal) y en Palabras de Francisco García son "
    "glosas palabra-por-palabra, no traducciones libres: útiles, pero deben "
    "ponderarse o filtrarse en el entrenamiento. La fuente web "
    "(lenguabribri.com) no pudo extraerse desde el entorno remoto por política de "
    "red (HTTP 403); su incorporación queda como mejora planificada.",
    size=9.5, color=GRIS)

para("5.3 Partición de datos", style="Heading 2")
para(
    "La partición es 80/10/10 estratificada por (dominio, confianza) con semilla "
    "42, resultando en 1 201 pares de entrenamiento, 152 de validación y 152 de "
    "prueba. La estratificación garantiza que cada split conserve la misma "
    "proporción de dominios y niveles de confianza, de modo que las métricas sean "
    "representativas del corpus completo. La trazabilidad por SHA-256 y el "
    "etiquetado por dominio/confianza permiten ablaciones reproducibles en fases "
    "posteriores.")

page_break()

# ============================================================== 6. MODELADO
para("6. Modelado: fine-tuning de modelos NMT", style="Heading 1")
para(
    "El modelado se desarrolló en iteraciones sucesivas (Avances 3 y 4), "
    "manteniendo constantes corpus, splits, semilla (42), longitud máxima, tamaño "
    "de lote y código de métricas para que las comparaciones fueran limpias. Se "
    "evaluó con tres métricas calculadas con sacrebleu: spBLEU (tokenización "
    "FLORES-200), chrF (F-score sobre n-gramas de caracteres, n=6) y chrF++ "
    "(chrF con n-gramas de palabras). En corpus pequeños y morfología rica, "
    "chrF/chrF++ son la señal de progreso más estable y se adoptan como métrica "
    "primaria.")

para("6.1 Avance 3 — Baseline NLLB-200 (Google Colab, T4)", style="Heading 2")
para(
    "Primera iteración: fine-tuning de facebook/nllb-200-distilled-600M en ambas "
    "direcciones (3 épocas, lr 5e-4). Como el bribri no tiene token de idioma "
    "propio en NLLB-200, se usó quy_Latn (quechua ayacucho) como proxy: lengua "
    "indígena americana en script latino presente en el modelo. La decisión es "
    "discutible y se documenta como tal.")
table(
    ["Dirección", "eval_loss ↓", "spBLEU ↑", "chrF ↑", "chrF++ ↑"],
    [["es → bri", "2.151", "18.61", "26.97", "26.01"],
     ["bri → es", "3.056", "10.26", "27.34", "26.16"],
     ["promedio", "2.603", "14.43", "27.16", "26.09"]],
    widths=[3.2, 3.3, 3.3, 3.3, 3.3],
    aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*4,
    highlight_row=2)
para(
    "Para un par ausente del preentrenamiento de NLLB, con ~1.2k pares y 3 "
    "épocas, chrF ≈ 27 en ambas direcciones es consistente con la literatura "
    "(cf. AmericasNLP, donde la línea base NLLB fine-tuneada para aymara y "
    "guaraní queda en chrF 25–35 con corpus de tamaño similar). La val loss cae "
    "monotónicamente sin rebote dentro del horizonte de 900 pasos: la pendiente "
    "final aún es negativa, señal de que el modelo no terminó de converger, "
    "observación que motivó directamente el Avance 4.")
figure(os.path.join(ROOT, "outputs", "training_curves.png"),
       "Figura 1. Curvas de entrenamiento del Avance 3 (val loss, spBLEU y "
       "chrF++ por pasos de optimización). chrF++ crece de forma estable; "
       "spBLEU mesetea con alta varianza por el reducido número de referencias.")

para("6.2 Avance 4 — Reentrenamiento en H100 (8 épocas, lr 2e-4)", style="Heading 2")
para(
    "Motivado por la no-convergencia del Avance 3, se reentrenó NLLB-200 en el "
    "servidor H100 de CENIA modificando únicamente épocas (3 → 8) y learning rate "
    "(5e-4 → 2e-4). Es el modelo seleccionado del proyecto.")
table(
    ["Dirección", "eval_loss ↓", "spBLEU ↑", "chrF ↑", "chrF++ ↑"],
    [["es → bri", "2.674", "28.74", "32.86", "32.35"],
     ["bri → es", "3.657", "13.57", "30.00", "28.58"],
     ["promedio", "3.166", "21.16", "31.43", "30.47"]],
    widths=[3.2, 3.3, 3.3, 3.3, 3.3],
    aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*4,
    highlight_row=2)
para(
    "La mejora respecto al Avance 3 es sustancial: chrF +15.7 % (27.16 → 31.43), "
    "chrF++ +16.8 % y spBLEU +46.6 %. Las épocas adicionales que el Avance 3 no "
    "ejecutó son las que producen la ganancia: chrF++ mesetea recién hacia el "
    "paso 1300–1500, mucho después de donde la val loss toca su mínimo.")
figure(os.path.join(ROOT, "outputs_nllb_h100", "training_curves_nllb_h100.png"),
       "Figura 2. Curvas de entrenamiento NLLB-H100 (24 evaluaciones, 100→2400 "
       "pasos). chrF++ sigue subiendo mucho después del mínimo de la val loss.")

para("6.3 El fenómeno del val loss creciente", style="Heading 2")
para(
    "Un resultado contraintuitivo y central: la corrida H100 tiene peor eval_loss "
    "(3.17 vs 2.60) y al mismo tiempo mejor chrF/spBLEU. No es un error, sino el "
    "comportamiento conocido del entrenamiento prolongado de modelos de "
    "generación. La cross-entropy de validación penaliza la confianza del modelo: "
    "al entrenar más épocas produce distribuciones más picudas (sobreconfiadas), "
    "de modo que cuando falla la penalización es mayor y la CE media sube. En "
    "cambio chrF/spBLEU miden el texto generado por decodificación, no la "
    "probabilidad token a token. La conclusión metodológica es que, en regímenes "
    "de bajo recurso, chrF/chrF++ deben gobernar la selección del modelo, no el "
    "val loss: seleccionar por val loss habría descartado el mejor modelo de "
    "traducción. Como efecto secundario, los bucles de repetición degenerada del "
    "Avance 3 desaparecen en la corrida H100.")

para("6.4 Baseline comparativo M2M-100 (418M)", style="Heading 2")
para(
    "Para validar empíricamente la elección de NLLB-200 se fine-tuneó M2M-100 "
    "(Meta) con idéntico pipeline, splits y métricas, y los mismos "
    "hiperparámetros del Avance 3. M2M-100 cubre 100 idiomas y no incluye ninguna "
    "lengua indígena americana; se usó br (bretón) como proxy. El resultado es "
    "drásticamente inferior y reincide en bucles de repetición masivos.")
table(
    ["Dirección", "eval_loss ↓", "spBLEU ↑", "chrF ↑", "chrF++ ↑"],
    [["es → bri", "2.901", "2.45", "10.49", "9.40"],
     ["bri → es", "4.348", "0.20", "8.27", "7.26"],
     ["promedio", "3.624", "1.33", "9.38", "8.33"]],
    widths=[3.2, 3.3, 3.3, 3.3, 3.3],
    aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*4,
    highlight_row=2)

para("6.5 Comparación de las tres corridas", style="Heading 2")
table(
    ["Modelo", "Configuración", "spBLEU ↑", "chrF ↑", "chrF++ ↑", "val_loss ↓"],
    [["NLLB-200 (H100)", "8 ep · lr 2e-4", "21.16", "31.43", "30.47", "3.17"],
     ["NLLB-200 (Colab)", "3 ep · lr 5e-4", "14.43", "27.16", "26.09", "2.60"],
     ["M2M-100", "3 ep · lr 5e-4", "1.33", "9.38", "8.33", "3.62"]],
    widths=[3.6, 3.0, 2.5, 2.4, 2.5, 2.5],
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT] +
           [WD_ALIGN_PARAGRAPH.CENTER]*4, highlight_row=0)
figure(os.path.join(ROOT, "outputs_nllb_h100", "comparison_bars_3way.png"),
       "Figura 3. Comparación de las tres corridas individuales. El baseline "
       "M2M-100 rinde muy por debajo (chrF 9.38 vs 31.43), confirmando que la "
       "cobertura lingüística de NLLB-200 es determinante para el bribri.",
       width_in=5.6)
para(
    "El baseline M2M-100 rinde drásticamente por debajo, confirmando "
    "empíricamente que la mayor cobertura lingüística de NLLB-200 —y en "
    "particular la disponibilidad de un proxy indígena americano real— es "
    "determinante. El modelo seleccionado para producción y ensamble es "
    "NLLB-200-distilled-600M reentrenado (8 épocas, lr 2e-4).")

page_break()

# ============================================================== 7. ENSAMBLES
para("7. Modelos de ensamble", style="Heading 1")
para(
    "La rúbrica de ensamble está formulada para clasificación/regresión (curva "
    "ROC, matriz de confusión, importancia de características). En NMT —tarea de "
    "generación de secuencias— esas herramientas no tienen análogo directo, pero "
    "las estrategias de ensamble sí existen y se documentan con sus equivalentes "
    "apropiados, evaluados con chrF/chrF++.")
para("Se implementaron dos estrategias:", space_after=2)
bullets([
    ("Ensamble homogéneo — checkpoint averaging.  ",
     "Promedia los pesos de los N mejores checkpoints de la corrida NLLB-H100. "
     "Es la analogía del bagging sobre el mismo modelo: reduce la varianza del "
     "punto final de entrenamiento sin coste de inferencia adicional (resulta un "
     "único modelo). Sólo es válido entre checkpoints del mismo vocabulario."),
    ("Ensamble heterogéneo — system combination por MBR.  ",
     "Como NLLB-200 y M2M-100 tienen vocabularios distintos, la combinación se "
     "hace a nivel de hipótesis: cada sistema genera su n-best, se juntan en un "
     "pool y se elige por sentencia la hipótesis de consenso (máxima chrF "
     "promedio contra el resto, Minimum Bayes Risk sin referencia). Es la "
     "analogía del voting/blending, anclada en el mejor modelo individual."),
])
para(
    "El objetivo de negocio del Programa Voces es preservación y revitalización, "
    "no traducción de producción a gran escala. En ese marco priman la calidad "
    "morfológica de la salida bribri (chrF), la estabilidad (ausencia de "
    "repeticiones) y un coste de inferencia bajo. El modelo individual NLLB-200 "
    "H100 ya domina las métricas y es de inferencia simple; el checkpoint "
    "averaging es el candidato natural de ensamble por no añadir coste de "
    "inferencia, mientras que el system combination heterogéneo, aunque puede "
    "subir chrF, duplica el coste de generación.")

page_break()

# ============================================================== 8. EVAL/DEPLOY
para("8. Evaluación y despliegue", style="Heading 1")
para(
    "Esta sección corresponde a la fase de Evaluación y Despliegue de "
    "CRISP-ML(Q). Su propósito es triple: determinar si el modelo es apto para "
    "implementarse o si es necesario retroceder a fases previas; exponer las "
    "decisiones y accionables derivados de los hallazgos, asignados a sus "
    "stakeholders; y proponer un entorno de producción fiable, escalable y "
    "eficiente.")

para("8.1 Criterios de éxito: formalización de la meta de Fase 0", style="Heading 2")
para(
    "En la Fase 0 el objetivo de negocio se definió cualitativamente "
    "—preservación lingüística asistida por tecnología— sin un umbral "
    "cuantitativo único. Esta ausencia es metodológicamente habitual y "
    "recomendable en escenarios de muy bajo recurso, donde fijar a priori un "
    "umbral tipo «chrF ≥ 50» sería arbitrario y desconectado del estado del arte "
    "para la lengua objetivo (Mager et al., 2018; Haddow et al., 2022). La fase "
    "de evaluación formaliza retroactivamente y de manera transparente el "
    "criterio de éxito, derivándolo del objetivo de negocio declarado, del estado "
    "del arte documentado (AmericasNLP; Ebrahimi et al., 2022) y de la naturaleza "
    "de prueba de concepto del Componente B.")
table(
    ["Nivel de criterio", "Definición operativa", "Umbral / evidencia"],
    [["Criterio mínimo (PoC aceptable)",
      "Valida que el enfoque (NLLB + fine-tuning) es viable para bribri y produce salida morfológicamente plausible, no degenerada, mejor que un baseline alternativo.",
      "chrF en el rango de la literatura (≈ 25–35); ausencia de bucles de repetición; superioridad sobre M2M-100.  ✓ Cumplido"],
     ["Criterio objetivo (utilidad asistida)",
      "Genera borradores útiles como insumo para un humano (lexicógrafo, docente, documentalista) que revisa y corrige, reduciendo su esfuerzo frente a partir de cero.",
      "Salida estructuralmente correcta en una proporción material de casos; léxico y morfología verificables.  ✓ Cumplido parcialmente"],
     ["Criterio NO pretendido (producción autónoma)",
      "Traducción confiable sin supervisión humana, apta para publicación directa o uso por usuario final.",
      "Requeriría chrF ≈ 50–60+ y control de alucinación.  ✗ Fuera de alcance declarado"]],
    widths=[3.6, 7.2, 5.7], font_size=9)
para(
    "Nota de honestidad metodológica. Se documenta explícitamente que el umbral "
    "cuantitativo no preexistía a esta fase. Construirlo aquí, anclado en el "
    "estado del arte y no en un número conveniente, es preferible a presentar un "
    "umbral inventado a posteriori que el modelo «casualmente» supere. La "
    "trazabilidad de esta decisión es, en sí misma, parte del criterio de calidad "
    "CRISP-ML(Q).", italic=True, color=GRIS, size=10)

para("8.2 ¿Es el rendimiento suficiente para producción?", style="Heading 2")
para(
    "La respuesta es condicionada al criterio de éxito. Para producción autónoma "
    "de propósito general: No. Un chrF de 31.43 está sustancialmente por debajo "
    "del rango de fluidez usable (chrF ≈ 50–60+ en pares de recurso medio; Goyal "
    "et al., 2022), y el análisis cualitativo evidenció tres patrones de error "
    "que prohíben el uso sin supervisión: traducción estructuralmente correcta "
    "pero parcial, alucinación temática (anclaje al dominio religioso mayoritario "
    "ante entradas ambiguas) y, en la corrida sub-entrenada, colapso a "
    "repetición.")
para(
    "Para producción restringida como prueba de concepto asistida: Sí. Bajo el "
    "criterio formalizado, el modelo cumple. Es confiable dentro de un dominio "
    "acotado y con validación humana en el bucle: asistencia a lexicógrafos, "
    "generación de borradores que un hablante revisa, apoyo pedagógico "
    "supervisado. Tras el reentrenamiento en H100 los bucles de repetición "
    "desaparecen, el modelo recupera estructuras y léxico en una proporción "
    "material de casos, y la dirección es→bri alcanza spBLEU 28.74, suficiente "
    "para que un revisor parta de un borrador y no de una página en blanco. La "
    "confiabilidad es condicional al flujo humano-en-el-bucle, no absoluta.")

para("8.3 ¿Hay margen de mejora?", style="Heading 2")
para(
    "Sí, y es amplio. El modelo no es un punto final, sino una base con palancas "
    "de mejora predecibles, ordenadas por relación esfuerzo/impacto:")
table(
    ["Palanca de mejora", "Fundamento", "Esfuerzo / impacto"],
    [["Ampliar el corpus",
      "1 201 pares están 1–2 órdenes de magnitud por debajo de lo típico. Incorporar la fuente web pendiente, aplicar OCR a fuentes escaneadas y digitalizar materiales del MEP elevaría el volumen.",
      "Alto / Muy alto"],
     ["Filtrado por confianza y data augmentation",
      "El 49 % del corpus es confidence=low (glosas). Ablaciones por confianza y back-translation mejorarían la señal sin nuevas fuentes.",
      "Medio / Medio-alto"],
     ["Optimización del proxy de idioma",
      "Se usó quy_Latn como proxy; alternativas incluyen extender el vocabulario con un token reasignado para bribri.",
      "Medio / Medio"],
     ["Ensamble (Avance 5)",
      "El checkpoint averaging reduce varianza sin coste extra; el MBR puede subir chrF duplicando el coste de generación.",
      "Bajo / Bajo-medio"],
     ["Scheduler y búsqueda de hiperparámetros",
      "El salto de +15.7 % al pasar de 3 a 8 épocas sugiere que la curva no está saturada; warmup y scheduler no se exploraron sistemáticamente.",
      "Bajo / Medio"]],
    widths=[3.8, 9.0, 3.7], font_size=9)
para(
    "La existencia de palancas de alto impacto aún sin agotar es, "
    "paradójicamente, un argumento a favor de implementar la PoC ahora: demuestra "
    "que la base actual no es un techo, sino un punto de partida con trayectoria "
    "de mejora identificada.")

para("8.4 Decisión: ¿implementar, iterar o retroceder?", style="Heading 2")
para(
    "La decisión fundamentada es no retroceder, sino implementar en producción "
    "restringida e iterar en paralelo:")
bullets([
    ("No se retrocede a modelado:  ",
     "el modelado cumplió su objetivo. NLLB-200 fue validado contra un baseline "
     "alternativo (M2M-100) y lo supera por 3.3×; la selección por chrF/chrF++ "
     "está justificada para el régimen de bajo recurso. No hay error que corregir."),
    ("No se retrocede a preparación de datos como bloqueante:  ",
     "el corpus v0 es limitado pero suficiente para una PoC, y su ampliación es "
     "una mejora futura planificada, no la corrección de un defecto."),
    ("Sí se implementa, con alcance restringido:  ",
     "la demo Streamlit ya operacionaliza el modelo; el despliegue entrega valor "
     "inmediato y genera retroalimentación —correcciones de hablantes— que "
     "alimenta la siguiente iteración del corpus."),
])
para("Recomendaciones clave de implementación:", bold=True, color=AZUL_OSC,
     space_after=2)
bullets([
    "Desplegar bajo encuadre explícito de PoC asistida: toda interfaz debe "
    "comunicar que las salidas son borradores que requieren validación humana.",
    "Mantener el human-in-the-loop como requisito de diseño, no como opción.",
    "Restringir el dominio de uso al documental/pedagógico, evitando usos de "
    "alto riesgo (legal, médico, oficial).",
    "Instrumentar la captura de correcciones de usuario (active learning de bajo coste).",
    "Gobernar los datos lingüísticos bajo principios CARE de soberanía de datos "
    "indígenas (Carroll et al., 2020), coordinando con ADITIBRI.",
    "Priorizar la ampliación del corpus como inversión de mayor retorno.",
])

para("8.5 Accionables por stakeholder", style="Heading 2")
para(
    "Tareas específicas asignadas al actor responsable, siguiendo el marco de "
    "Miller (2022) para la identificación de stakeholders en proyectos de IA.")
table(
    ["#", "Accionable", "Stakeholder responsable", "Entregable / cierre", "Horizonte"],
    [["A1", "Desplegar la demo Streamlit en cloud gestionado con encuadre de PoC y aviso de validación humana.",
      "Equipo CR — ingeniería", "URL pública con el checkpoint NLLB-H100 cargado.", "Corto"],
     ["A2", "Validar cualitativamente una muestra estratificada con hablantes nativos.",
      "ADITIBRI + hablantes, coord. D. Leiva", "Protocolo + reporte de aceptabilidad por dominio.", "Corto-medio"],
     ["A3", "Incorporar la fuente web pendiente y aplicar OCR a fuentes escaneadas para ampliar el corpus.",
      "Equipo CR — datos", "Corpus v1 con volumen ≥ 2× y reporte actualizado.", "Medio"],
     ["A4", "Proveer presupuesto de inferencia y reentrenamiento (H100) para escalamiento.",
      "CENIA — C. Aspillaga (sponsor)", "Asignación de cómputo y línea presupuestaria.", "Medio"],
     ["A5", "Instrumentar la captura de correcciones de usuario como insumo de active learning.",
      "Equipo CR — ingeniería", "Log estructurado (origen / borrador / corrección).", "Medio"],
     ["A6", "Establecer el marco de gobernanza de datos bajo principios CARE.",
      "CENIA + PUC + ADITIBRI; superv. C. Villaseñor", "Acuerdo de gobernanza y licenciamiento del corpus.", "Medio-largo"],
     ["A7", "Validar la pertinencia científica de las decisiones metodológicas de cara a publicación.",
      "Asesoría académica (C. Villaseñor) + Tec", "Visto bueno metodológico / borrador de artículo.", "Largo"]],
    widths=[0.8, 6.0, 3.6, 4.3, 1.8], font_size=8.5,
    aligns=[WD_ALIGN_PARAGRAPH.CENTER] + [WD_ALIGN_PARAGRAPH.LEFT]*4)

para("8.6 Análisis de plataformas cloud", style="Heading 2")
para(
    "Se evalúan los cuatro proveedores señalados por la actividad —AWS, "
    "Microsoft Azure, Google Cloud (GCP) e IBM watsonx— frente a los requisitos "
    "de la solución: servir un modelo Transformer de ~2.4 GB para inferencia, con "
    "capacidad puntual de reentrenamiento en GPU, a coste contenido y con baja "
    "fricción operativa para un equipo pequeño. Los factores de evaluación son "
    "facilidad de uso, servicios de ML/NMT, escalabilidad, costos/capa académica "
    "y continuidad con la infraestructura existente.")
table(
    ["Factor", "AWS", "Microsoft Azure", "Google Cloud", "IBM watsonx"],
    [["Servicio ML principal", "SageMaker (suite madura)", "Azure ML + AI Foundry", "Vertex AI (GenAI/MLOps)", "watsonx.ai (empresarial)"],
     ["GPU bajo demanda", "Amplio (A100, H100)", "Amplio (H100, ND)", "A100/H100 y TPUs", "Disponible, acotado"],
     ["Integración Hugging Face", "Muy buena (DLC)", "Buena (HF en AML)", "Muy buena (Vertex+HF)", "Limitada (Granite)"],
     ["Facilidad de uso", "Curva pronunciada", "Buena, fluida", "Buena, moderna", "Media, orientada a TI"],
     ["Escalabilidad", "Excelente", "Excelente", "Excelente", "Buena"],
     ["Costos / capa académica", "Free tier; Educate", "Crédito + for Students", "Free tier; for Education", "Tier Lite acotado"],
     ["Continuidad con el proyecto", "Baja (migración)", "Alta (H100 ya en Azure)", "Media (exp. previa)", "Baja"]],
    widths=[3.4, 3.3, 3.4, 3.2, 3.2], font_size=8.5,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*4)
para(
    "Proveedor recomendado: Microsoft Azure, con Google Cloud (Vertex AI) como "
    "alternativa secundaria. La elección no se sostiene en superioridad técnica "
    "absoluta —los cuatro hiperescaladores son comparables para servir un modelo "
    "de este tamaño— sino en el factor que minimiza riesgo y coste: la "
    "continuidad de infraestructura. El reentrenamiento que produjo el modelo se "
    "ejecutó en un servidor H100 provisto por CENIA en Azure; desplegar la "
    "inferencia en el mismo proveedor elimina migración de pesos y credenciales, "
    "reutiliza la relación institucional y mantiene datos y cómputo bajo una sola "
    "gobernanza, relevante para los principios de soberanía de datos (A6). Se "
    "descarta AWS por mayor coste de migración sin beneficio diferencial, e IBM "
    "watsonx por su orientación a modelos propios (Granite) y menor afinidad con "
    "el stack Hugging Face / NLLB.")

para("8.7 Arquitectura de producción propuesta", style="Heading 2")
para("Arquitectura de despliegue mínima viable, escalable por capas:",
     space_after=2)
bullets([
    ("Capa de servicio:  ", "la aplicación Streamlit (app.py) contenedurizada, "
     "sirviendo el checkpoint NLLB-200 fine-tuneado, como endpoint gestionado en "
     "Azure Machine Learning (o App Service con contenedor para la PoC inicial)."),
    ("Capa de inferencia:  ", "instancia con GPU bajo demanda para baja latencia, "
     "con fallback a CPU para contener costos; no_repeat_ngram_size=3 activo."),
    ("Capa de datos y registro:  ", "almacenamiento de objetos para los pesos "
     "(~2.4 GB) y registro estructurado de correcciones de usuario (A5)."),
    ("Capa de gobernanza:  ", "control de acceso y versionado del corpus bajo el "
     "acuerdo de soberanía de datos (A6)."),
    ("Confiabilidad y eficiencia:  ", "escalado a cero en inactividad para uso "
     "intermitente; monitoreo de latencia y de tasa de salidas degeneradas como "
     "señal de regresión."),
])

page_break()

# ============================================================== 9. CONCLUSIONES
para("9. Conclusiones generales", style="Heading 1")
para(
    "El sistema NMT bribri–español alcanza, en su mejor configuración "
    "(NLLB-200-distilled-600M, H100, 8 épocas), chrF 31.43 / chrF++ 30.47 / "
    "spBLEU 21.16. Frente a los criterios de éxito formalizados, el modelo no es "
    "apto para traducción autónoma de propósito general, pero sí cumple como "
    "prueba de concepto funcional para asistencia documental y pedagógica con "
    "validación humana en el bucle. La decisión fundamentada es implementar en "
    "producción restringida e iterar en paralelo, sin retroceder a fases "
    "anteriores: el modelado cumplió su objetivo y la limitación del corpus es "
    "una mejora planificada, no un defecto a corregir.")
para("El valor del entregable para el Programa Voces es triple:", space_after=2)
bullets([
    "Valida empíricamente que el enfoque NLLB + fine-tuning es el camino correcto "
    "para bribri, descartando alternativas (M2M-100) con evidencia.",
    "Entrega una base reproducible, trazable y medible sobre la cual CENIA puede "
    "escalar, con palancas de mejora ya identificadas y priorizadas.",
    "Materializa una herramienta de asistencia que hoy no existe para la "
    "comunidad bribri, generando además, mediante su uso, los datos de corrección "
    "que alimentarán la siguiente iteración.",
])
para(
    "La plataforma de despliegue recomendada es Microsoft Azure, por continuidad "
    "con la infraestructura de cómputo del proyecto, con Google Cloud (Vertex AI) "
    "como alternativa secundaria. Más allá del resultado métrico, el principal "
    "aporte metodológico del proyecto es haber demostrado, con trazabilidad "
    "completa, que en regímenes de muy bajo recurso la selección de modelo debe "
    "gobernarse por métricas robustas a morfología (chrF/chrF++) y no por la "
    "cross-entropy de validación, y que un criterio de éxito honesto, anclado en "
    "el estado del arte, es preferible a un umbral inventado a posteriori.")

para("9.1 Limitaciones y trabajo futuro", style="Heading 2")
bullets([
    "Tamaño del corpus: 1 201 pares de entrenamiento están 1–2 órdenes de "
    "magnitud por debajo de lo típico para fine-tuning estable de NLLB-200.",
    "Ruido de glosas: el 49 % del corpus es confidence=low; no se filtró para no "
    "reducir el set a ~600 pares.",
    "Proxy de idioma: quy_Latn se usó como ancla; extender el vocabulario con un "
    "token propio de bribri queda como exploración pendiente.",
    "Fuente web no incorporada por política de red; su extracción ampliaría el corpus.",
    "Trabajo futuro: ampliación del corpus (web + OCR + MEP), validación con "
    "hablantes nativos, back-translation, búsqueda sistemática de hiperparámetros "
    "y publicación de los resultados metodológicos.",
])

page_break()

# ============================================================== REFERENCIAS
para("10. Referencias", style="Heading 1")
refs = [
    "Amazon Web Services. (2024). Amazon SageMaker: Machine learning service. https://aws.amazon.com/sagemaker/",
    "Carroll, S. R., Garba, I., Figueroa-Rodríguez, O. L., Holbrook, J., Lovett, R., Materechera, S., Parsons, M., Raseroka, K., Rodriguez-Lonebear, D., Rowe, R., Sara, R., Walker, J. D., Anderson, J., & Hudson, M. (2020). The CARE principles for indigenous data governance. Data Science Journal, 19(1), 43. https://doi.org/10.5334/dsj-2020-043",
    "Ebrahimi, A., Mager, M., Oncevay, A., Chaudhary, V., Chiruzzo, L., Fan, A., Ortega, J., Ramos, R., Rios, A., Vladimir, M., Giménez-Lugo, G. A., Mager, E., Neubig, G., Palmer, A., Coto-Solano, R., Vu, N. T., & Kann, K. (2022). AmericasNLI: Evaluating zero-shot natural language understanding of pretrained multilingual models in truly low-resource languages. Proceedings of the 60th Annual Meeting of the ACL, 6279–6299. https://doi.org/10.18653/v1/2022.acl-long.435",
    "Goyal, N., Gao, C., Chaudhary, V., Chen, P.-J., Wenzek, G., Ju, D., Krishnan, S., Ranzato, M., Guzmán, F., & Fan, A. (2022). The Flores-101 evaluation benchmark for low-resource and multilingual machine translation. Transactions of the ACL, 10, 522–538. https://doi.org/10.1162/tacl_a_00474",
    "Google Cloud. (2024). Vertex AI: Unified platform for machine learning. https://cloud.google.com/vertex-ai",
    "Haddow, B., Bawden, R., Miceli Barone, A. V., Helcl, J., & Birch, A. (2022). Survey of low-resource machine translation. Computational Linguistics, 48(3), 673–732. https://doi.org/10.1162/coli_a_00446",
    "IBM. (2024). watsonx.ai: Enterprise studio for AI builders. https://www.ibm.com/products/watsonx-ai",
    "INEC. (2011). X Censo Nacional de Población y VI de Vivienda 2011. Instituto Nacional de Estadística y Censos de Costa Rica.",
    "Jara Murillo, C. V. (2018). El bribri: Lengua en peligro. Revista de Filología y Lingüística de la Universidad de Costa Rica, 44(1), 15–32.",
    "Jara Murillo, C. V., & García Segura, A. (1997). Se'ttö': Hablemos bribri. Editorial de la Universidad de Costa Rica.",
    "Mager, M., Gutierrez-Vasques, X., Sierra, G., & Meza-Ruiz, I. (2018). Challenges of language technologies for the indigenous languages of the Americas. Proceedings of the 27th International Conference on Computational Linguistics, 55–69.",
    "Margery Peña, E. (1989). Diccionario fraseológico bribri-español, español-bribri. Editorial de la Universidad de Costa Rica.",
    "Microsoft. (2024). Azure Machine Learning: Enterprise-grade ML service. https://azure.microsoft.com/products/machine-learning",
    "Miller, G. J. (2022). Stakeholder roles in artificial intelligence projects. Project Leadership and Society, 3, 100068. https://doi.org/10.1016/j.plas.2022.100068",
    "NLLB Team, Costa-jussà, M. R., Cross, J., Çelebi, O., Elbayad, M., Heafield, K., et al. (2022). No language left behind: Scaling human-centered machine translation. arXiv. https://doi.org/10.48550/arXiv.2207.04672",
    "Ortega, J., Maldonado, A., & Mager, M. (2020). Neural machine translation for low-resource indigenous languages. Proceedings of AmericasNLP.",
    "Popović, M. (2015). chrF: Character n-gram F-score for automatic MT evaluation. Proceedings of the Tenth Workshop on Statistical Machine Translation, 392–395. https://doi.org/10.18653/v1/W15-3049",
    "Studer, S., Bui, T. B., Drescher, C., Hanuschkin, A., Winkler, L., Peters, S., & Müller, K.-R. (2021). Towards CRISP-ML(Q): A machine learning process model with quality assurance methodology. Machine Learning and Knowledge Extraction, 3(2), 392–413. https://doi.org/10.3390/make3020020",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, A. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30, 5998–6008.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(r)
    run.font.size = Pt(9.5)

# ----------------------------------------------------- pie de página numerado
def add_footer_pagenums(document):
    for section in document.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        run = fp.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
        it.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(it); run._r.append(f2)
        run.font.size = Pt(9); run.font.color.rgb = GRIS

add_footer_pagenums(doc)

doc.save(OUT)
print("Documento guardado en:", OUT)
